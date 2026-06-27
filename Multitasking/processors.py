"""
Document, PDF, and image processing functions.
Self-contained module — no Streamlit dependency, so it can be tested
and reused independently of the UI layer.
"""
import io
import os
import html
import subprocess

from pypdf import PdfReader, PdfWriter
import pikepdf
from PIL import Image, ImageDraw, ImageFont, ImageOps, ImageFilter
from pdf2image import convert_from_bytes
from pdf2docx import Converter as PDF2DocxConverter
import docx as docx_lib
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import qrcode

WORKDIR = "/tmp/linguaai_toolkit"
os.makedirs(WORKDIR, exist_ok=True)

# ════════════════════════════════════════════════
#  PDF TOOLS
# ════════════════════════════════════════════════

def pdf_merge(file_bytes_list: list) -> bytes:
    writer = PdfWriter()
    for fb in file_bytes_list:
        reader = PdfReader(io.BytesIO(fb))
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_split(file_bytes: bytes) -> list:
    """Returns list of (filename, bytes) for each page."""
    reader = PdfReader(io.BytesIO(file_bytes))
    results = []
    for i, page in enumerate(reader.pages):
        writer = PdfWriter()
        writer.add_page(page)
        out = io.BytesIO()
        writer.write(out)
        results.append((f"page_{i+1}.pdf", out.getvalue()))
    return results


def pdf_split_range(file_bytes: bytes, start: int, end: int) -> bytes:
    """Extract a page range (1-indexed, inclusive) into one PDF."""
    reader = PdfReader(io.BytesIO(file_bytes))
    writer = PdfWriter()
    for i in range(start - 1, min(end, len(reader.pages))):
        writer.add_page(reader.pages[i])
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_rotate(file_bytes: bytes, degrees: int, pages: list = None) -> bytes:
    reader = PdfReader(io.BytesIO(file_bytes))
    writer = PdfWriter()
    for i, page in enumerate(reader.pages):
        if pages is None or (i + 1) in pages:
            page.rotate(degrees)
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_compress(file_bytes: bytes, image_quality: int = 60) -> bytes:
    """
    Compress a PDF by re-saving with object stream compression and
    down-sampling embedded images via pikepdf + Pillow.
    """
    in_path = os.path.join(WORKDIR, "_compress_in.pdf")
    out_path = os.path.join(WORKDIR, "_compress_out.pdf")
    with open(in_path, "wb") as f:
        f.write(file_bytes)

    with pikepdf.open(in_path) as pdf:
        for page in pdf.pages:
            for img_key in list(page.images.keys()):
                try:
                    raw_img = page.images[img_key]
                    pdfimg = pikepdf.PdfImage(raw_img)
                    pil_img = pdfimg.as_pil_image()
                    if pil_img.mode in ("RGBA", "P"):
                        pil_img = pil_img.convert("RGB")
                    buf = io.BytesIO()
                    pil_img.save(buf, format="JPEG", quality=image_quality, optimize=True)
                    buf.seek(0)
                    raw_img.write(buf.read(), filter=pikepdf.Name("/DCTDecode"))
                except Exception:
                    continue
        pdf.save(out_path, compress_streams=True,
                 object_stream_mode=pikepdf.ObjectStreamMode.generate)

    with open(out_path, "rb") as f:
        return f.read()


def pdf_compress_simple(file_bytes: bytes) -> bytes:
    """Lighter-weight compression: strips metadata, optimizes object streams."""
    in_path = os.path.join(WORKDIR, "_c_in.pdf")
    out_path = os.path.join(WORKDIR, "_c_out.pdf")
    with open(in_path, "wb") as f:
        f.write(file_bytes)
    with pikepdf.open(in_path) as pdf:
        pdf.save(out_path, compress_streams=True,
                  object_stream_mode=pikepdf.ObjectStreamMode.generate,
                  linearize=False)
    with open(out_path, "rb") as f:
        return f.read()


def pdf_add_watermark(file_bytes: bytes, watermark_text: str, opacity: float = 0.3) -> bytes:
    """Add a diagonal text watermark to every page."""
    reader = PdfReader(io.BytesIO(file_bytes))
    page0 = reader.pages[0]
    width = float(page0.mediabox.width)
    height = float(page0.mediabox.height)

    wm_buf = io.BytesIO()
    c = pdf_canvas.Canvas(wm_buf, pagesize=(width, height))
    c.saveState()
    c.setFont("Helvetica-Bold", min(width, height) / 10)
    c.setFillColorRGB(0.6, 0.6, 0.6, alpha=opacity)
    c.translate(width / 2, height / 2)
    c.rotate(45)
    c.drawCentredString(0, 0, watermark_text)
    c.restoreState()
    c.save()
    wm_buf.seek(0)
    wm_reader = PdfReader(wm_buf)
    wm_page = wm_reader.pages[0]

    writer = PdfWriter()
    for page in reader.pages:
        page.merge_page(wm_page)
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_protect(file_bytes: bytes, user_password: str, owner_password: str = None) -> bytes:
    reader = PdfReader(io.BytesIO(file_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(user_password, owner_password or user_password)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_unlock(file_bytes: bytes, password: str) -> bytes:
    reader = PdfReader(io.BytesIO(file_bytes), password=password)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def pdf_extract_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for i, page in enumerate(reader.pages):
        text_parts.append(f"--- Page {i+1} ---\n{page.extract_text()}")
    return "\n\n".join(text_parts)


def pdf_extract_images(file_bytes: bytes) -> list:
    """Returns list of (filename, bytes) for extracted images."""
    in_path = os.path.join(WORKDIR, "_extract_in.pdf")
    with open(in_path, "wb") as f:
        f.write(file_bytes)
    results = []
    with pikepdf.open(in_path) as pdf:
        for pg_num, page in enumerate(pdf.pages):
            for img_num, (key, raw_img) in enumerate(page.images.items()):
                try:
                    pdfimg = pikepdf.PdfImage(raw_img)
                    pil_img = pdfimg.as_pil_image()
                    buf = io.BytesIO()
                    fmt = "PNG"
                    if pil_img.mode in ("RGBA", "P"):
                        pil_img = pil_img.convert("RGB")
                    pil_img.save(buf, format=fmt)
                    results.append((f"page{pg_num+1}_img{img_num+1}.png", buf.getvalue()))
                except Exception:
                    continue
    return results


def pdf_to_images(file_bytes: bytes, dpi: int = 150) -> list:
    """Convert each PDF page to a PNG image. Returns list of (filename, bytes)."""
    images = convert_from_bytes(file_bytes, dpi=dpi)
    results = []
    for i, img in enumerate(images):
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        results.append((f"page_{i+1}.png", buf.getvalue()))
    return results


def images_to_pdf(image_bytes_list: list) -> bytes:
    """Combine multiple images into a single PDF."""
    images = []
    for img_bytes in image_bytes_list:
        img = Image.open(io.BytesIO(img_bytes))
        if img.mode != "RGB":
            img = img.convert("RGB")
        images.append(img)
    if not images:
        return b""
    out = io.BytesIO()
    images[0].save(out, format="PDF", save_all=True, append_images=images[1:])
    return out.getvalue()


def pdf_get_page_count(file_bytes: bytes) -> int:
    return len(PdfReader(io.BytesIO(file_bytes)).pages)


# ════════════════════════════════════════════════
#  WORD ↔ PDF CONVERSION
# ════════════════════════════════════════════════

def pdf_to_word(file_bytes: bytes) -> bytes:
    """Convert PDF to a .docx file using pdf2docx (layout-preserving)."""
    in_path = os.path.join(WORKDIR, "_p2w_in.pdf")
    out_path = os.path.join(WORKDIR, "_p2w_out.docx")
    with open(in_path, "wb") as f:
        f.write(file_bytes)
    cv = PDF2DocxConverter(in_path)
    cv.convert(out_path)
    cv.close()
    with open(out_path, "rb") as f:
        return f.read()


def word_to_pdf(file_bytes: bytes) -> bytes:
    """
    Convert .docx to PDF. Tries LibreOffice headless first (best fidelity);
    falls back to a pure-Python text-extraction + reportlab rebuild if
    LibreOffice isn't available on the host.
    """
    in_path = os.path.join(WORKDIR, "_w2p_in.docx")
    out_dir = WORKDIR
    with open(in_path, "wb") as f:
        f.write(file_bytes)

    # Try LibreOffice headless conversion (best fidelity, preserves formatting)
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, in_path],
            capture_output=True, timeout=60
        )
        out_path = os.path.join(out_dir, "_w2p_in.pdf")
        if result.returncode == 0 and os.path.exists(out_path):
            with open(out_path, "rb") as f:
                return f.read()
    except Exception:
        pass

    # Fallback: extract text/paragraphs from docx and rebuild as a simple PDF
    return _word_to_pdf_fallback(file_bytes)


def _word_to_pdf_fallback(file_bytes: bytes) -> bytes:
    """Pure-Python fallback: rebuilds basic text content as a PDF (no LibreOffice)."""
    doc = docx_lib.Document(io.BytesIO(file_bytes))
    out = io.BytesIO()
    pdf_doc = SimpleDocTemplate(out, pagesize=letter,
                                  topMargin=0.8*inch, bottomMargin=0.8*inch,
                                  leftMargin=0.8*inch, rightMargin=0.8*inch)
    styles = getSampleStyleSheet()
    story = []
    for para in doc.paragraphs:
        text = html.escape(para.text)
        if not text.strip():
            story.append(Spacer(1, 8))
            continue
        style_name = "Normal"
        if para.style and para.style.name:
            if "Heading 1" in para.style.name or "Title" in para.style.name:
                style_name = "Heading1"
            elif "Heading 2" in para.style.name:
                style_name = "Heading2"
            elif "Heading" in para.style.name:
                style_name = "Heading3"
        story.append(Paragraph(text, styles[style_name]))
        story.append(Spacer(1, 4))
    if not story:
        story = [Paragraph("(Empty document)", styles["Normal"])]
    pdf_doc.build(story)
    return out.getvalue()


def text_to_docx(text: str, title: str = "Document") -> bytes:
    doc = docx_lib.Document()
    if title:
        h = doc.add_heading(title, level=1)
    for para_text in text.split("\n"):
        if para_text.strip():
            doc.add_paragraph(para_text)
        else:
            doc.add_paragraph("")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def text_to_pdf(text: str, title: str = "Document") -> bytes:
    out = io.BytesIO()
    pdf_doc = SimpleDocTemplate(out, pagesize=letter,
                                  topMargin=0.8*inch, bottomMargin=0.8*inch,
                                  leftMargin=0.8*inch, rightMargin=0.8*inch)
    styles = getSampleStyleSheet()
    story = []
    if title:
        story.append(Paragraph(html.escape(title), styles["Title"]))
        story.append(Spacer(1, 12))
    for para_text in text.split("\n"):
        if para_text.strip():
            story.append(Paragraph(html.escape(para_text), styles["Normal"]))
            story.append(Spacer(1, 6))
        else:
            story.append(Spacer(1, 6))
    pdf_doc.build(story)
    return out.getvalue()


def docx_extract_text(file_bytes: bytes) -> str:
    doc = docx_lib.Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


# ════════════════════════════════════════════════
#  IMAGE TOOLS
# ════════════════════════════════════════════════

def image_convert_format(file_bytes: bytes, target_format: str, quality: int = 90) -> bytes:
    img = Image.open(io.BytesIO(file_bytes))
    out = io.BytesIO()
    fmt = target_format.upper()
    if fmt in ("JPEG", "JPG") and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    save_fmt = "JPEG" if fmt == "JPG" else fmt
    save_kwargs = {}
    if save_fmt in ("JPEG", "WEBP"):
        save_kwargs["quality"] = quality
    img.save(out, format=save_fmt, **save_kwargs)
    return out.getvalue()


def image_compress(file_bytes: bytes, quality: int = 60, max_dimension: int = None) -> bytes:
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    if max_dimension:
        ratio = min(max_dimension / img.width, max_dimension / img.height, 1.0)
        if ratio < 1.0:
            img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=quality, optimize=True)
    return out.getvalue()


def image_resize(file_bytes: bytes, width: int, height: int, keep_aspect: bool = True) -> bytes:
    img = Image.open(io.BytesIO(file_bytes))
    if keep_aspect:
        img.thumbnail((width, height), Image.LANCZOS)
    else:
        img = img.resize((width, height), Image.LANCZOS)
    out = io.BytesIO()
    fmt = img.format or "PNG"
    if img.mode == "RGBA" and fmt == "JPEG":
        img = img.convert("RGB")
    img.save(out, format=fmt)
    return out.getvalue()


def image_rotate_flip(file_bytes: bytes, rotate_deg: int = 0, flip_h: bool = False, flip_v: bool = False) -> bytes:
    img = Image.open(io.BytesIO(file_bytes))
    if rotate_deg:
        img = img.rotate(-rotate_deg, expand=True)
    if flip_h:
        img = ImageOps.mirror(img)
    if flip_v:
        img = ImageOps.flip(img)
    out = io.BytesIO()
    img.save(out, format=img.format or "PNG")
    return out.getvalue()


def image_add_watermark_text(file_bytes: bytes, text: str, opacity: int = 128, position: str = "center") -> bytes:
    img = Image.open(io.BytesIO(file_bytes)).convert("RGBA")
    overlay = Image.new("RGBA", img.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    font_size = max(20, img.width // 15)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", font_size)
    except Exception:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    positions = {
        "center": ((img.width - tw) // 2, (img.height - th) // 2),
        "bottom-right": (img.width - tw - 20, img.height - th - 20),
        "top-left": (20, 20),
    }
    pos = positions.get(position, positions["center"])
    draw.text(pos, text, font=font, fill=(255, 255, 255, opacity))
    combined = Image.alpha_composite(img, overlay).convert("RGB")
    out = io.BytesIO()
    combined.save(out, format="PNG")
    return out.getvalue()


def image_to_grayscale(file_bytes: bytes) -> bytes:
    img = Image.open(io.BytesIO(file_bytes)).convert("L")
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def image_apply_filter(file_bytes: bytes, filter_name: str) -> bytes:
    img = Image.open(io.BytesIO(file_bytes))
    filters = {
        "Blur": ImageFilter.BLUR,
        "Sharpen": ImageFilter.SHARPEN,
        "Edge Enhance": ImageFilter.EDGE_ENHANCE,
        "Contour": ImageFilter.CONTOUR,
        "Emboss": ImageFilter.EMBOSS,
        "Smooth": ImageFilter.SMOOTH,
    }
    if filter_name in filters:
        img = img.filter(filters[filter_name])
    out = io.BytesIO()
    img.save(out, format=img.format or "PNG")
    return out.getvalue()


def text_to_image(text: str, width: int = 800, height: int = 400,
                    bg_color: str = "#0f1420", text_color: str = "#e2e8f8",
                    font_size: int = 32) -> bytes:
    """Render plain text onto an image canvas (simple text-to-image)."""
    img = Image.new("RGB", (width, height), bg_color)
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", font_size)
    except Exception:
        font = ImageFont.load_default()

    # Simple word-wrap
    words = text.split()
    lines, current = [], ""
    for word in words:
        test = (current + " " + word).strip()
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] > width - 60 and current:
            lines.append(current)
            current = word
        else:
            current = test
    if current:
        lines.append(current)

    line_height = font_size + 10
    total_h = len(lines) * line_height
    y = max(20, (height - total_h) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = (width - (bbox[2] - bbox[0])) // 2
        draw.text((x, y), line, font=font, fill=text_color)
        y += line_height

    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def generate_qr_code(data: str, box_size: int = 10, fill_color: str = "black", back_color: str = "white") -> bytes:
    qr = qrcode.QRCode(version=1, box_size=box_size, border=4)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color=fill_color, back_color=back_color)
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


# ════════════════════════════════════════════════
#  PDF FORM / DATA TOOLS
# ════════════════════════════════════════════════

def pdf_to_text_file(file_bytes: bytes) -> bytes:
    return pdf_extract_text(file_bytes).encode("utf-8")
